#!/usr/bin/env python3
"""
종합적인 문서 수집 파이프라인 테스트 스크립트

이 스크립트는 문서 수집 파이프라인의 모든 단계를 테스트합니다:
1. 파일 포맷 지원 테스트 (PDF, TXT, MD, DOCX, HTML, CSV, JSON)
2. 텍스트 추출 및 청킹
3. 메타데이터 추출
4. PII 필터링
5. 임베딩 생성
6. 데이터베이스 저장
7. 오류 처리
8. 배치 처리 성능

모든 테스트는 실제 파일을 생성하고 처리하며, 한글 처리와 성능 메트릭을 포함합니다.
"""

import asyncio
import logging
import os
import sys
import json
import time
import traceback
import shutil
import psutil
from pathlib import Path
from datetime import datetime, timedelta
from typing import Dict, List, Any, Optional, Tuple
import tempfile
import zipfile
from dataclasses import dataclass, field
from collections import defaultdict
import csv
import io

# 테스트용 파일 생성 라이브러리
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import docx
from docx.shared import Inches

# 한글 처리를 위한 인코딩 설정
import locale
import codecs

# 프로젝트 모듈
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from apps.ingestion.document_parser import get_parser_factory, ParsedDocument
from apps.ingestion.chunking_strategy import default_chunker, DocumentChunk
from apps.ingestion.pii_filter import default_pii_filter, PIIFilterResult, PIIType, MaskingStrategy
from apps.ingestion.ingestion_pipeline import IngestionPipeline, ProcessingResult, ProcessingStatus
from apps.api.database import db_manager, EmbeddingService, Document, DocumentChunk as DBChunk, Embedding

# 로깅 설정
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('document_ingestion_test.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

@dataclass
class TestMetrics:
    """테스트 메트릭"""
    start_time: datetime = field(default_factory=datetime.now)
    end_time: Optional[datetime] = None
    total_files: int = 0
    successful_files: int = 0
    failed_files: int = 0
    total_chunks: int = 0
    total_characters: int = 0
    pii_detections: int = 0
    embeddings_generated: int = 0
    processing_times: Dict[str, List[float]] = field(default_factory=lambda: defaultdict(list))
    memory_usage: List[float] = field(default_factory=list)
    errors: List[Dict[str, Any]] = field(default_factory=list)

    @property
    def duration(self) -> float:
        if self.end_time:
            return (self.end_time - self.start_time).total_seconds()
        return (datetime.now() - self.start_time).total_seconds()

    @property
    def success_rate(self) -> float:
        if self.total_files == 0:
            return 0.0
        return (self.successful_files / self.total_files) * 100

    @property
    def average_processing_time(self) -> float:
        all_times = []
        for times in self.processing_times.values():
            all_times.extend(times)
        return sum(all_times) / len(all_times) if all_times else 0.0

class DocumentIngestionTester:
    """문서 수집 파이프라인 종합 테스터"""

    def __init__(self):
        self.temp_dir = Path(tempfile.mkdtemp(prefix="ingestion_test_"))
        self.metrics = TestMetrics()
        self.parser_factory = get_parser_factory()
        self.chunker = default_chunker
        self.pii_filter = default_pii_filter
        self.embedding_service = EmbeddingService()
        self.pipeline = IngestionPipeline()

        logger.info(f"테스트 임시 디렉토리: {self.temp_dir}")

    def __del__(self):
        """정리"""
        if hasattr(self, 'temp_dir') and self.temp_dir.exists():
            shutil.rmtree(self.temp_dir, ignore_errors=True)

    def monitor_memory(self):
        """메모리 사용량 모니터링"""
        process = psutil.Process()
        memory_mb = process.memory_info().rss / 1024 / 1024
        self.metrics.memory_usage.append(memory_mb)
        return memory_mb

    async def create_test_files(self) -> Dict[str, Path]:
        """다양한 포맷의 테스트 파일 생성"""
        logger.info("테스트 파일 생성 중...")
        files = {}

        # 1. 텍스트 파일 (한글 포함)
        txt_content = """
안녕하세요! 이것은 한글 텍스트 파일입니다.
개인정보 테스트:
- 이메일: hong@example.com
- 전화번호: 010-1234-5678
- 주민등록번호: 123456-1234567
- 신용카드: 1234-5678-9012-3456

기술 문서 내용:
Dynamic Taxonomy RAG 시스템은 계층적 문서 분류와 검색을 제공합니다.
이 시스템은 자동으로 문서를 분류하고 관련 정보를 검색할 수 있습니다.

Lorem ipsum dolor sit amet, consectetur adipiscing elit.
Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua.
        """.strip()

        txt_file = self.temp_dir / "test_korean.txt"
        with open(txt_file, 'w', encoding='utf-8') as f:
            f.write(txt_content)
        files['txt'] = txt_file

        # 2. Markdown 파일
        md_content = """
# 한글 마크다운 문서
## 개요
이것은 **한글 마크다운** 테스트 파일입니다.

### 개인정보 섹션
- 사용자: 김철수
- 이메일: kim.cs@company.co.kr
- 연락처: 02-123-4567
- 카드번호: 5678-1234-9012-3456

### 기술 내용
```python
def process_document(text):
    return chunker.chunk(text)
```

> 이것은 인용문입니다.

### 목록
1. 첫 번째 항목
2. 두 번째 항목
3. 세 번째 항목

[링크 예시](https://example.com)
        """.strip()

        md_file = self.temp_dir / "test_korean.md"
        with open(md_file, 'w', encoding='utf-8') as f:
            f.write(md_content)
        files['md'] = md_file

        # 3. HTML 파일
        html_content = """
<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="UTF-8">
    <title>한글 HTML 테스트</title>
</head>
<body>
    <h1>한글 HTML 문서</h1>
    <p>이것은 한글 HTML 테스트 파일입니다.</p>

    <h2>개인정보</h2>
    <ul>
        <li>이름: 박영희</li>
        <li>이메일: park@test.com</li>
        <li>전화: 010-9876-5432</li>
        <li>주민번호: 987654-2345678</li>
    </ul>

    <h2>기술 내용</h2>
    <pre><code>
    async def embed_text(text):
        return await embedding_service.embed(text)
    </code></pre>

    <script>
        console.log("스크립트는 제거됩니다");
    </script>
</body>
</html>
        """.strip()

        html_file = self.temp_dir / "test_korean.html"
        with open(html_file, 'w', encoding='utf-8') as f:
            f.write(html_content)
        files['html'] = html_file

        # 4. CSV 파일
        csv_content = [
            ['이름', '이메일', '전화번호', '부서'],
            ['김민수', 'kim@company.com', '010-1111-2222', '개발팀'],
            ['이영희', 'lee@company.com', '010-3333-4444', '마케팅팀'],
            ['박철수', 'park@company.com', '010-5555-6666', '영업팀'],
            ['개인정보', 'private@secret.com', '010-7777-8888', '기밀부서']
        ]

        csv_file = self.temp_dir / "test_korean.csv"
        with open(csv_file, 'w', encoding='utf-8', newline='') as f:
            writer = csv.writer(f)
            writer.writerows(csv_content)
        files['csv'] = csv_file

        # 5. JSON 파일
        json_content = {
            "title": "한글 JSON 문서",
            "description": "JSON 포맷 테스트",
            "users": [
                {
                    "name": "홍길동",
                    "email": "hong@example.org",
                    "phone": "010-1234-0000",
                    "ssn": "123456-1111111"
                },
                {
                    "name": "김영수",
                    "email": "kim@test.org",
                    "phone": "010-5678-0000",
                    "credit_card": "1111-2222-3333-4444"
                }
            ],
            "content": {
                "korean": "한글 콘텐츠입니다.",
                "english": "This is English content.",
                "technical": "RAG 시스템의 임베딩 벡터는 1536차원입니다."
            }
        }

        json_file = self.temp_dir / "test_korean.json"
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump(json_content, f, ensure_ascii=False, indent=2)
        files['json'] = json_file

        # 6. PDF 파일 생성 (reportlab 사용)
        await self._create_pdf_file(files)

        # 7. DOCX 파일 생성
        await self._create_docx_file(files)

        # 8. 손상된 파일 (오류 처리 테스트용)
        corrupted_file = self.temp_dir / "corrupted.pdf"
        with open(corrupted_file, 'wb') as f:
            f.write(b"This is not a valid PDF file content")
        files['corrupted'] = corrupted_file

        # 9. 빈 파일
        empty_file = self.temp_dir / "empty.txt"
        empty_file.touch()
        files['empty'] = empty_file

        # 10. 대용량 파일 (성능 테스트용)
        large_content = "한글 대용량 테스트 콘텐츠입니다. " * 10000
        large_file = self.temp_dir / "large_test.txt"
        with open(large_file, 'w', encoding='utf-8') as f:
            f.write(large_content)
        files['large'] = large_file

        logger.info(f"총 {len(files)}개의 테스트 파일 생성 완료")
        return files

    async def _create_pdf_file(self, files: Dict[str, Path]):
        """PDF 파일 생성"""
        try:
            # 한글 폰트 등록 시도
            try:
                # Windows 시스템 폰트 사용
                pdfmetrics.registerFont(TTFont('Malgun', 'malgun.ttf'))
                font_name = 'Malgun'
            except:
                try:
                    # 대체 폰트
                    pdfmetrics.registerFont(TTFont('Nanum', 'NanumGothic.ttf'))
                    font_name = 'Nanum'
                except:
                    # 기본 폰트 사용 (한글 깨짐 가능)
                    font_name = 'Helvetica'

            pdf_file = self.temp_dir / "test_korean.pdf"
            c = canvas.Canvas(str(pdf_file), pagesize=letter)

            # PDF 메타데이터 설정
            c.setTitle("한글 PDF 테스트 문서")
            c.setAuthor("테스트 작성자")
            c.setSubject("PDF 파싱 테스트")

            y = 750
            c.setFont(font_name, 16)
            c.drawString(50, y, "한글 PDF 테스트 문서")

            y -= 30
            c.setFont(font_name, 12)
            content_lines = [
                "이것은 PDF 파싱 테스트용 문서입니다.",
                "",
                "개인정보 테스트:",
                "- 이메일: pdf@test.com",
                "- 전화: 010-1111-0000",
                "- 주민번호: 111111-2222222",
                "",
                "기술 내용:",
                "PDF 문서는 복잡한 레이아웃을 가질 수 있습니다.",
                "텍스트 추출 시 정확성이 중요합니다.",
                "",
                "영어 내용:",
                "This is English content in the PDF.",
                "Mixed language support is important.",
            ]

            for line in content_lines:
                if y < 50:  # 새 페이지
                    c.showPage()
                    y = 750
                    c.setFont(font_name, 12)

                try:
                    c.drawString(50, y, line)
                except:
                    # 한글 폰트 실패 시 영어만
                    if line.encode('ascii', errors='ignore').decode('ascii'):
                        c.drawString(50, y, line.encode('ascii', errors='ignore').decode('ascii'))
                y -= 20

            c.save()
            files['pdf'] = pdf_file

        except Exception as e:
            logger.warning(f"PDF 파일 생성 실패: {e}")
            # 간단한 텍스트 기반 PDF 생성
            pdf_file = self.temp_dir / "test_simple.pdf"
            c = canvas.Canvas(str(pdf_file), pagesize=letter)
            c.setFont('Helvetica', 12)
            c.drawString(50, 750, "Simple PDF Test Document")
            c.drawString(50, 730, "Email: simple@test.com")
            c.drawString(50, 710, "Phone: 010-0000-0000")
            c.save()
            files['pdf'] = pdf_file

    async def _create_docx_file(self, files: Dict[str, Path]):
        """DOCX 파일 생성"""
        try:
            docx_file = self.temp_dir / "test_korean.docx"
            doc = docx.Document()

            # 메타데이터 설정
            core_props = doc.core_properties
            core_props.title = "한글 DOCX 테스트 문서"
            core_props.author = "테스트 작성자"
            core_props.subject = "DOCX 파싱 테스트"
            core_props.created = datetime.now()

            # 제목
            title = doc.add_heading('한글 DOCX 테스트 문서', 0)

            # 내용
            doc.add_paragraph('이것은 DOCX 파싱 테스트용 문서입니다.')

            # 개인정보 섹션
            doc.add_heading('개인정보 테스트', level=1)
            pii_para = doc.add_paragraph()
            pii_para.add_run('이메일: ').bold = True
            pii_para.add_run('docx@test.com')
            pii_para.add_run('\n전화번호: ').bold = True
            pii_para.add_run('010-2222-0000')
            pii_para.add_run('\n주민번호: ').bold = True
            pii_para.add_run('222222-3333333')

            # 기술 내용
            doc.add_heading('기술 내용', level=1)
            doc.add_paragraph('DOCX 문서는 구조화된 콘텐츠를 포함할 수 있습니다.')
            doc.add_paragraph('표, 이미지, 스타일 등을 지원합니다.')

            # 표 추가
            table = doc.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '이름'
            hdr_cells[1].text = '이메일'
            hdr_cells[2].text = '부서'

            row_cells = table.add_row().cells
            row_cells[0].text = '김테스트'
            row_cells[1].text = 'test@docx.com'
            row_cells[2].text = '개발팀'

            doc.save(str(docx_file))
            files['docx'] = docx_file

        except Exception as e:
            logger.warning(f"DOCX 파일 생성 실패: {e}")

    async def test_file_format_support(self, files: Dict[str, Path]) -> Dict[str, Any]:
        """파일 포맷 지원 테스트"""
        logger.info("=== 파일 포맷 지원 테스트 시작 ===")
        results = {}

        for fmt, file_path in files.items():
            if fmt in ['corrupted', 'empty']:
                continue

            start_time = time.time()
            self.monitor_memory()

            try:
                parser = self.parser_factory.get_parser(file_path)
                if not parser:
                    results[fmt] = {
                        'success': False,
                        'error': 'No parser available',
                        'processing_time': 0
                    }
                    continue

                parsed_doc = await parser.parse(file_path)
                processing_time = time.time() - start_time

                results[fmt] = {
                    'success': True,
                    'text_length': len(parsed_doc.content),
                    'metadata_keys': list(parsed_doc.metadata.keys()),
                    'processing_time': processing_time,
                    'has_korean': any(ord(c) > 127 for c in parsed_doc.content[:100])
                }

                self.metrics.processing_times['parsing'].append(processing_time)
                self.metrics.total_characters += len(parsed_doc.content)

                logger.info(f"{fmt} 파싱 성공: {len(parsed_doc.content)} 문자, {processing_time:.3f}초")

            except Exception as e:
                results[fmt] = {
                    'success': False,
                    'error': str(e),
                    'processing_time': time.time() - start_time
                }
                logger.error(f"{fmt} 파싱 실패: {e}")
                self.metrics.errors.append({
                    'stage': 'parsing',
                    'file_format': fmt,
                    'error': str(e),
                    'traceback': traceback.format_exc()
                })

        return results

    async def test_text_extraction_and_chunking(self, files: Dict[str, Path]) -> Dict[str, Any]:
        """텍스트 추출 및 청킹 테스트"""
        logger.info("=== 텍스트 추출 및 청킹 테스트 시작 ===")
        results = {}

        for fmt, file_path in files.items():
            if fmt in ['corrupted', 'empty']:
                continue

            start_time = time.time()
            self.monitor_memory()

            try:
                # 파싱
                parser = self.parser_factory.get_parser(file_path)
                if not parser:
                    continue

                parsed_doc = await parser.parse(file_path)

                # 청킹
                chunks = await self.chunker.chunk_document(parsed_doc)
                processing_time = time.time() - start_time

                results[fmt] = {
                    'success': True,
                    'original_length': len(parsed_doc.content),
                    'chunk_count': len(chunks),
                    'chunk_sizes': [len(chunk.content) for chunk in chunks],
                    'average_chunk_size': sum(len(chunk.content) for chunk in chunks) / len(chunks) if chunks else 0,
                    'processing_time': processing_time,
                    'overlap_check': self._check_chunk_overlap(chunks)
                }

                self.metrics.processing_times['chunking'].append(processing_time)
                self.metrics.total_chunks += len(chunks)

                logger.info(f"{fmt} 청킹 성공: {len(chunks)}개 청크, {processing_time:.3f}초")

            except Exception as e:
                results[fmt] = {
                    'success': False,
                    'error': str(e),
                    'processing_time': time.time() - start_time
                }
                logger.error(f"{fmt} 청킹 실패: {e}")
                self.metrics.errors.append({
                    'stage': 'chunking',
                    'file_format': fmt,
                    'error': str(e),
                    'traceback': traceback.format_exc()
                })

        return results

    def _check_chunk_overlap(self, chunks: List[DocumentChunk]) -> Dict[str, Any]:
        """청크 오버랩 검사"""
        if len(chunks) < 2:
            return {'has_overlap': False, 'overlap_count': 0}

        overlap_count = 0
        for i in range(len(chunks) - 1):
            current_end = chunks[i].content[-50:]  # 마지막 50문자
            next_start = chunks[i + 1].content[:50]  # 처음 50문자

            # 간단한 오버랩 체크
            if any(word in next_start for word in current_end.split() if len(word) > 3):
                overlap_count += 1

        return {
            'has_overlap': overlap_count > 0,
            'overlap_count': overlap_count,
            'overlap_ratio': overlap_count / (len(chunks) - 1) if len(chunks) > 1 else 0
        }

    async def test_metadata_extraction(self, files: Dict[str, Path]) -> Dict[str, Any]:
        """메타데이터 추출 테스트"""
        logger.info("=== 메타데이터 추출 테스트 시작 ===")
        results = {}

        for fmt, file_path in files.items():
            if fmt in ['corrupted', 'empty']:
                continue

            start_time = time.time()

            try:
                parser = self.parser_factory.get_parser(file_path)
                if not parser:
                    continue

                parsed_doc = await parser.parse(file_path)
                processing_time = time.time() - start_time

                # 메타데이터 검증
                metadata = parsed_doc.metadata
                required_fields = ['file_name', 'file_size', 'created_at']
                optional_fields = ['title', 'author', 'subject', 'creator', 'modified_at']

                results[fmt] = {
                    'success': True,
                    'metadata': metadata,
                    'required_fields_present': {field: field in metadata for field in required_fields},
                    'optional_fields_present': {field: field in metadata for field in optional_fields},
                    'total_fields': len(metadata),
                    'processing_time': processing_time
                }

                logger.info(f"{fmt} 메타데이터 추출 성공: {len(metadata)}개 필드")

            except Exception as e:
                results[fmt] = {
                    'success': False,
                    'error': str(e),
                    'processing_time': time.time() - start_time
                }
                logger.error(f"{fmt} 메타데이터 추출 실패: {e}")

        return results

    async def test_pii_filtering(self, files: Dict[str, Path]) -> Dict[str, Any]:
        """PII 필터링 테스트"""
        logger.info("=== PII 필터링 테스트 시작 ===")
        results = {}

        # 테스트할 PII 패턴
        test_patterns = {
            'email': ['hong@example.com', 'kim.cs@company.co.kr'],
            'phone': ['010-1234-5678', '02-123-4567'],
            'ssn': ['123456-1234567', '987654-2345678'],
            'credit_card': ['1234-5678-9012-3456', '5678-1234-9012-3456']
        }

        for fmt, file_path in files.items():
            if fmt in ['corrupted', 'empty']:
                continue

            start_time = time.time()
            self.monitor_memory()

            try:
                parser = self.parser_factory.get_parser(file_path)
                if not parser:
                    continue

                parsed_doc = await parser.parse(file_path)

                # PII 필터링 테스트
                filter_result = await self.pii_filter.filter_text(
                    parsed_doc.content,
                    masking_strategy=MaskingStrategy.REDACT
                )

                processing_time = time.time() - start_time

                # PII 탐지 정확성 검사
                detected_types = [detection.pii_type for detection in filter_result.detections]

                results[fmt] = {
                    'success': True,
                    'original_length': len(parsed_doc.content),
                    'filtered_length': len(filter_result.filtered_text),
                    'detections_count': len(filter_result.detections),
                    'detected_types': detected_types,
                    'detections': [
                        {
                            'type': d.pii_type.value,
                            'confidence': d.confidence,
                            'start': d.start_pos,
                            'end': d.end_pos,
                            'original': d.original_text,
                            'masked': d.masked_text
                        } for d in filter_result.detections
                    ],
                    'processing_time': processing_time,
                    'accuracy_check': self._check_pii_accuracy(parsed_doc.content, filter_result, test_patterns)
                }

                self.metrics.processing_times['pii_filtering'].append(processing_time)
                self.metrics.pii_detections += len(filter_result.detections)

                logger.info(f"{fmt} PII 필터링 성공: {len(filter_result.detections)}개 탐지, {processing_time:.3f}초")

            except Exception as e:
                results[fmt] = {
                    'success': False,
                    'error': str(e),
                    'processing_time': time.time() - start_time
                }
                logger.error(f"{fmt} PII 필터링 실패: {e}")
                self.metrics.errors.append({
                    'stage': 'pii_filtering',
                    'file_format': fmt,
                    'error': str(e),
                    'traceback': traceback.format_exc()
                })

        return results

    def _check_pii_accuracy(self, original_text: str, filter_result: PIIFilterResult,
                           test_patterns: Dict[str, List[str]]) -> Dict[str, Any]:
        """PII 탐지 정확성 검사"""
        accuracy = {}

        for pii_type, patterns in test_patterns.items():
            detected_count = 0
            total_count = 0

            for pattern in patterns:
                total_count += original_text.count(pattern)

                # 해당 패턴이 탐지되었는지 확인
                for detection in filter_result.detections:
                    if pattern in detection.original_text:
                        detected_count += 1
                        break

            if total_count > 0:
                accuracy[pii_type] = {
                    'detected': detected_count,
                    'total': total_count,
                    'accuracy': detected_count / total_count
                }

        return accuracy

    async def test_embedding_generation(self, files: Dict[str, Path]) -> Dict[str, Any]:
        """임베딩 생성 테스트"""
        logger.info("=== 임베딩 생성 테스트 시작 ===")
        results = {}

        for fmt, file_path in files.items():
            if fmt in ['corrupted', 'empty', 'large']:  # 대용량 파일은 임베딩 테스트에서 제외
                continue

            start_time = time.time()
            self.monitor_memory()

            try:
                parser = self.parser_factory.get_parser(file_path)
                if not parser:
                    continue

                parsed_doc = await parser.parse(file_path)
                chunks = await self.chunker.chunk_document(parsed_doc)

                # 임베딩 생성 (최대 3개 청크로 제한)
                test_chunks = chunks[:3]
                embeddings = []

                for chunk in test_chunks:
                    try:
                        embedding = await self.embedding_service.embed(chunk.content)
                        embeddings.append({
                            'dimension': len(embedding),
                            'norm': sum(x*x for x in embedding) ** 0.5,
                            'sample_values': embedding[:5]  # 처음 5개 값만
                        })
                    except Exception as e:
                        logger.warning(f"임베딩 생성 실패 (청크): {e}")
                        embeddings.append({'error': str(e)})

                processing_time = time.time() - start_time

                results[fmt] = {
                    'success': True,
                    'chunk_count': len(chunks),
                    'tested_chunks': len(test_chunks),
                    'successful_embeddings': len([e for e in embeddings if 'error' not in e]),
                    'embeddings': embeddings,
                    'processing_time': processing_time
                }

                self.metrics.processing_times['embedding'].append(processing_time)
                self.metrics.embeddings_generated += len([e for e in embeddings if 'error' not in e])

                logger.info(f"{fmt} 임베딩 생성 성공: {len([e for e in embeddings if 'error' not in e])}/{len(test_chunks)}개")

            except Exception as e:
                results[fmt] = {
                    'success': False,
                    'error': str(e),
                    'processing_time': time.time() - start_time
                }
                logger.error(f"{fmt} 임베딩 생성 실패: {e}")
                self.metrics.errors.append({
                    'stage': 'embedding',
                    'file_format': fmt,
                    'error': str(e),
                    'traceback': traceback.format_exc()
                })

        return results

    async def test_database_storage(self, files: Dict[str, Path]) -> Dict[str, Any]:
        """데이터베이스 저장 테스트"""
        logger.info("=== 데이터베이스 저장 테스트 시작 ===")
        results = {}

        # 테스트용 파일 하나만 선택 (TXT)
        test_file = files.get('txt')
        if not test_file:
            return {'error': 'No test file available'}

        start_time = time.time()
        self.monitor_memory()

        try:
            # 전체 파이프라인 실행
            processing_result = await self.pipeline.process_document(test_file)
            processing_time = time.time() - start_time

            results['full_pipeline'] = {
                'success': processing_result.status == ProcessingStatus.COMPLETED,
                'status': processing_result.status.value,
                'document_id': processing_result.document_id,
                'chunk_count': len(processing_result.chunks) if processing_result.chunks else 0,
                'embedding_count': len(processing_result.embeddings) if processing_result.embeddings else 0,
                'processing_time': processing_time,
                'error': processing_result.error
            }

            # 데이터베이스 검증
            if processing_result.status == ProcessingStatus.COMPLETED and processing_result.document_id:
                verification = await self._verify_database_storage(processing_result.document_id)
                results['database_verification'] = verification

            self.metrics.processing_times['database_storage'].append(processing_time)

            logger.info(f"데이터베이스 저장 테스트 완료: {processing_result.status.value}")

        except Exception as e:
            results['full_pipeline'] = {
                'success': False,
                'error': str(e),
                'processing_time': time.time() - start_time
            }
            logger.error(f"데이터베이스 저장 테스트 실패: {e}")
            self.metrics.errors.append({
                'stage': 'database_storage',
                'error': str(e),
                'traceback': traceback.format_exc()
            })

        return results

    async def _verify_database_storage(self, document_id: str) -> Dict[str, Any]:
        """데이터베이스 저장 검증"""
        try:
            async with db_manager.get_session() as session:
                # 문서 확인
                doc = await session.get(Document, document_id)
                if not doc:
                    return {'success': False, 'error': 'Document not found in database'}

                # 청크 확인
                from sqlalchemy import select
                chunk_stmt = select(DBChunk).where(DBChunk.document_id == document_id)
                chunk_result = await session.execute(chunk_stmt)
                chunks = chunk_result.scalars().all()

                # 임베딩 확인
                embedding_stmt = select(Embedding).where(Embedding.document_id == document_id)
                embedding_result = await session.execute(embedding_stmt)
                embeddings = embedding_result.scalars().all()

                return {
                    'success': True,
                    'document_exists': True,
                    'document_title': doc.title,
                    'chunk_count': len(chunks),
                    'embedding_count': len(embeddings),
                    'metadata_keys': list(doc.metadata.keys()) if doc.metadata else []
                }

        except Exception as e:
            return {'success': False, 'error': str(e)}

    async def test_error_handling(self, files: Dict[str, Path]) -> Dict[str, Any]:
        """오류 처리 테스트"""
        logger.info("=== 오류 처리 테스트 시작 ===")
        results = {}

        # 1. 손상된 파일 테스트
        corrupted_file = files.get('corrupted')
        if corrupted_file:
            start_time = time.time()
            try:
                processing_result = await self.pipeline.process_document(corrupted_file)
                results['corrupted_file'] = {
                    'handled_gracefully': processing_result.status == ProcessingStatus.FAILED,
                    'status': processing_result.status.value,
                    'error_message': processing_result.error,
                    'processing_time': time.time() - start_time
                }
            except Exception as e:
                results['corrupted_file'] = {
                    'handled_gracefully': False,
                    'unhandled_exception': str(e),
                    'processing_time': time.time() - start_time
                }

        # 2. 빈 파일 테스트
        empty_file = files.get('empty')
        if empty_file:
            start_time = time.time()
            try:
                processing_result = await self.pipeline.process_document(empty_file)
                results['empty_file'] = {
                    'handled_gracefully': processing_result.status in [ProcessingStatus.FAILED, ProcessingStatus.SKIPPED],
                    'status': processing_result.status.value,
                    'error_message': processing_result.error,
                    'processing_time': time.time() - start_time
                }
            except Exception as e:
                results['empty_file'] = {
                    'handled_gracefully': False,
                    'unhandled_exception': str(e),
                    'processing_time': time.time() - start_time
                }

        # 3. 존재하지 않는 파일 테스트
        nonexistent_file = self.temp_dir / "nonexistent.txt"
        start_time = time.time()
        try:
            processing_result = await self.pipeline.process_document(nonexistent_file)
            results['nonexistent_file'] = {
                'handled_gracefully': processing_result.status == ProcessingStatus.FAILED,
                'status': processing_result.status.value,
                'error_message': processing_result.error,
                'processing_time': time.time() - start_time
            }
        except Exception as e:
            results['nonexistent_file'] = {
                'handled_gracefully': False,
                'unhandled_exception': str(e),
                'processing_time': time.time() - start_time
            }

        return results

    async def test_batch_processing(self, files: Dict[str, Path]) -> Dict[str, Any]:
        """배치 처리 성능 테스트"""
        logger.info("=== 배치 처리 성능 테스트 시작 ===")

        # 정상적인 파일들만 선택
        valid_files = [files[fmt] for fmt in ['txt', 'md', 'html', 'csv', 'json'] if fmt in files]

        if not valid_files:
            return {'error': 'No valid files for batch processing test'}

        start_time = time.time()
        initial_memory = self.monitor_memory()

        try:
            # 순차 처리
            sequential_start = time.time()
            sequential_results = []
            for file_path in valid_files:
                result = await self.pipeline.process_document(file_path)
                sequential_results.append(result)
            sequential_time = time.time() - sequential_start

            # 메모리 사용량 체크
            peak_memory = max(self.metrics.memory_usage[-10:]) if self.metrics.memory_usage else initial_memory

            # 병렬 처리 (제한된 동시성)
            parallel_start = time.time()
            semaphore = asyncio.Semaphore(3)  # 최대 3개 동시 처리

            async def process_with_semaphore(file_path):
                async with semaphore:
                    return await self.pipeline.process_document(file_path)

            parallel_tasks = [process_with_semaphore(file_path) for file_path in valid_files]
            parallel_results = await asyncio.gather(*parallel_tasks, return_exceptions=True)
            parallel_time = time.time() - parallel_start

            # 결과 분석
            successful_sequential = sum(1 for r in sequential_results if r.status == ProcessingStatus.COMPLETED)
            successful_parallel = sum(1 for r in parallel_results if not isinstance(r, Exception) and r.status == ProcessingStatus.COMPLETED)

            total_time = time.time() - start_time

            results = {
                'success': True,
                'total_files': len(valid_files),
                'sequential_processing': {
                    'time': sequential_time,
                    'successful': successful_sequential,
                    'avg_time_per_file': sequential_time / len(valid_files)
                },
                'parallel_processing': {
                    'time': parallel_time,
                    'successful': successful_parallel,
                    'avg_time_per_file': parallel_time / len(valid_files),
                    'speedup_ratio': sequential_time / parallel_time if parallel_time > 0 else 0
                },
                'memory_usage': {
                    'initial_mb': initial_memory,
                    'peak_mb': peak_memory,
                    'increase_mb': peak_memory - initial_memory
                },
                'total_time': total_time
            }

            logger.info(f"배치 처리 완료: 순차 {sequential_time:.3f}초, 병렬 {parallel_time:.3f}초")

        except Exception as e:
            results = {
                'success': False,
                'error': str(e),
                'total_time': time.time() - start_time
            }
            logger.error(f"배치 처리 테스트 실패: {e}")

        return results

    def finalize_metrics(self):
        """메트릭 최종화"""
        self.metrics.end_time = datetime.now()

        # 성공률 계산
        total_files = len([f for fmt, f in self.__dict__.get('test_files', {}).items() if fmt not in ['corrupted', 'empty']])
        successful_files = total_files - len([e for e in self.metrics.errors if e['stage'] != 'expected_failure'])

        self.metrics.total_files = total_files
        self.metrics.successful_files = successful_files
        self.metrics.failed_files = total_files - successful_files

    def generate_report(self, all_results: Dict[str, Any]) -> Dict[str, Any]:
        """종합 테스트 보고서 생성"""
        self.finalize_metrics()

        report = {
            'summary': {
                'test_start_time': self.metrics.start_time.isoformat(),
                'test_end_time': self.metrics.end_time.isoformat(),
                'total_duration_seconds': self.metrics.duration,
                'total_files_tested': self.metrics.total_files,
                'successful_files': self.metrics.successful_files,
                'failed_files': self.metrics.failed_files,
                'success_rate_percent': self.metrics.success_rate,
                'total_chunks_created': self.metrics.total_chunks,
                'total_characters_processed': self.metrics.total_characters,
                'pii_detections_count': self.metrics.pii_detections,
                'embeddings_generated': self.metrics.embeddings_generated
            },
            'performance_metrics': {
                'average_processing_time': self.metrics.average_processing_time,
                'processing_times_by_stage': {
                    stage: {
                        'count': len(times),
                        'average': sum(times) / len(times) if times else 0,
                        'min': min(times) if times else 0,
                        'max': max(times) if times else 0
                    } for stage, times in self.metrics.processing_times.items()
                },
                'memory_usage': {
                    'samples': len(self.metrics.memory_usage),
                    'average_mb': sum(self.metrics.memory_usage) / len(self.metrics.memory_usage) if self.metrics.memory_usage else 0,
                    'peak_mb': max(self.metrics.memory_usage) if self.metrics.memory_usage else 0,
                    'min_mb': min(self.metrics.memory_usage) if self.metrics.memory_usage else 0
                }
            },
            'detailed_results': all_results,
            'error_analysis': {
                'total_errors': len(self.metrics.errors),
                'errors_by_stage': {},
                'errors_by_format': {},
                'error_details': self.metrics.errors[:10]  # 처음 10개 오류만
            },
            'recommendations': self._generate_recommendations(all_results)
        }

        # 오류 분석
        for error in self.metrics.errors:
            stage = error.get('stage', 'unknown')
            fmt = error.get('file_format', 'unknown')

            if stage not in report['error_analysis']['errors_by_stage']:
                report['error_analysis']['errors_by_stage'][stage] = 0
            report['error_analysis']['errors_by_stage'][stage] += 1

            if fmt not in report['error_analysis']['errors_by_format']:
                report['error_analysis']['errors_by_format'][fmt] = 0
            report['error_analysis']['errors_by_format'][fmt] += 1

        return report

    def _generate_recommendations(self, all_results: Dict[str, Any]) -> List[str]:
        """개선 권장사항 생성"""
        recommendations = []

        # 성공률 기반 권장사항
        if self.metrics.success_rate < 90:
            recommendations.append(f"전체 성공률이 {self.metrics.success_rate:.1f}%로 낮습니다. 오류 처리 개선이 필요합니다.")

        # 성능 기반 권장사항
        avg_time = self.metrics.average_processing_time
        if avg_time > 5.0:
            recommendations.append(f"평균 처리 시간이 {avg_time:.2f}초로 느립니다. 성능 최적화를 고려하세요.")

        # 메모리 사용량 기반 권장사항
        if self.metrics.memory_usage:
            peak_memory = max(self.metrics.memory_usage)
            if peak_memory > 500:  # 500MB 이상
                recommendations.append(f"최대 메모리 사용량이 {peak_memory:.1f}MB입니다. 메모리 최적화가 필요할 수 있습니다.")

        # PII 탐지 기반 권장사항
        pii_results = all_results.get('pii_filtering', {})
        low_accuracy_formats = []
        for fmt, result in pii_results.items():
            if result.get('success') and 'accuracy_check' in result:
                accuracy = result['accuracy_check']
                for pii_type, acc_data in accuracy.items():
                    if acc_data.get('accuracy', 1.0) < 0.8:
                        low_accuracy_formats.append(f"{fmt}-{pii_type}")

        if low_accuracy_formats:
            recommendations.append(f"PII 탐지 정확도가 낮은 형태: {', '.join(low_accuracy_formats)}. PII 필터 개선이 필요합니다.")

        # 청킹 기반 권장사항
        chunking_results = all_results.get('text_extraction_and_chunking', {})
        large_chunks = []
        for fmt, result in chunking_results.items():
            if result.get('success') and result.get('average_chunk_size', 0) > 2000:
                large_chunks.append(fmt)

        if large_chunks:
            recommendations.append(f"청크 크기가 큰 형태: {', '.join(large_chunks)}. 청킹 전략 조정을 고려하세요.")

        if not recommendations:
            recommendations.append("모든 테스트가 양호하게 통과했습니다. 현재 설정을 유지하세요.")

        return recommendations

    async def run_all_tests(self) -> Dict[str, Any]:
        """모든 테스트 실행"""
        logger.info("🚀 종합적인 문서 수집 파이프라인 테스트 시작")

        all_results = {}

        try:
            # 1. 테스트 파일 생성
            logger.info("1️⃣ 테스트 파일 생성")
            test_files = await self.create_test_files()
            self.test_files = test_files
            all_results['file_creation'] = {
                'success': True,
                'files_created': len(test_files),
                'file_types': list(test_files.keys())
            }

            # 2. 파일 포맷 지원 테스트
            logger.info("2️⃣ 파일 포맷 지원 테스트")
            all_results['file_format_support'] = await self.test_file_format_support(test_files)

            # 3. 텍스트 추출 및 청킹 테스트
            logger.info("3️⃣ 텍스트 추출 및 청킹 테스트")
            all_results['text_extraction_and_chunking'] = await self.test_text_extraction_and_chunking(test_files)

            # 4. 메타데이터 추출 테스트
            logger.info("4️⃣ 메타데이터 추출 테스트")
            all_results['metadata_extraction'] = await self.test_metadata_extraction(test_files)

            # 5. PII 필터링 테스트
            logger.info("5️⃣ PII 필터링 테스트")
            all_results['pii_filtering'] = await self.test_pii_filtering(test_files)

            # 6. 임베딩 생성 테스트
            logger.info("6️⃣ 임베딩 생성 테스트")
            all_results['embedding_generation'] = await self.test_embedding_generation(test_files)

            # 7. 데이터베이스 저장 테스트
            logger.info("7️⃣ 데이터베이스 저장 테스트")
            all_results['database_storage'] = await self.test_database_storage(test_files)

            # 8. 오류 처리 테스트
            logger.info("8️⃣ 오류 처리 테스트")
            all_results['error_handling'] = await self.test_error_handling(test_files)

            # 9. 배치 처리 성능 테스트
            logger.info("9️⃣ 배치 처리 성능 테스트")
            all_results['batch_processing'] = await self.test_batch_processing(test_files)

            logger.info("✅ 모든 테스트 완료")

        except Exception as e:
            logger.error(f"테스트 실행 중 오류: {e}")
            all_results['fatal_error'] = {
                'error': str(e),
                'traceback': traceback.format_exc()
            }

        return all_results

async def main():
    """메인 테스트 함수"""
    tester = DocumentIngestionTester()

    try:
        # 모든 테스트 실행
        results = await tester.run_all_tests()

        # 보고서 생성
        report = tester.generate_report(results)

        # 결과 저장
        report_file = Path("document_ingestion_test_report.json")
        with open(report_file, 'w', encoding='utf-8') as f:
            json.dump(report, f, ensure_ascii=False, indent=2)

        # 요약 출력
        print("\n" + "="*80)
        print("📊 문서 수집 파이프라인 테스트 결과 요약")
        print("="*80)

        summary = report['summary']
        print(f"🕐 테스트 시간: {summary['total_duration_seconds']:.2f}초")
        print(f"📁 처리된 파일: {summary['total_files_tested']}개")
        print(f"✅ 성공: {summary['successful_files']}개")
        print(f"❌ 실패: {summary['failed_files']}개")
        print(f"📈 성공률: {summary['success_rate_percent']:.1f}%")
        print(f"📄 생성된 청크: {summary['total_chunks_created']}개")
        print(f"🔤 처리된 문자: {summary['total_characters_processed']:,}개")
        print(f"🛡️ PII 탐지: {summary['pii_detections_count']}개")
        print(f"🧠 임베딩 생성: {summary['embeddings_generated']}개")

        print("\n⚡ 성능 메트릭:")
        perf = report['performance_metrics']
        print(f"평균 처리 시간: {perf['average_processing_time']:.3f}초")
        print(f"최대 메모리 사용량: {perf['memory_usage']['peak_mb']:.1f}MB")

        if report['recommendations']:
            print("\n💡 권장사항:")
            for i, rec in enumerate(report['recommendations'], 1):
                print(f"{i}. {rec}")

        print(f"\n📋 상세 보고서: {report_file.absolute()}")
        print("="*80)

        # 성공 여부 반환
        success_rate = summary['success_rate_percent']
        if success_rate >= 90:
            print("🎉 테스트 성공! 파이프라인이 정상적으로 작동합니다.")
            return 0
        elif success_rate >= 70:
            print("⚠️ 테스트 부분 성공. 일부 개선이 필요합니다.")
            return 1
        else:
            print("❌ 테스트 실패. 심각한 문제가 있습니다.")
            return 2

    except Exception as e:
        logger.error(f"테스트 실행 실패: {e}")
        print(f"\n💥 테스트 실행 중 치명적 오류 발생: {e}")
        return 3

    finally:
        # 정리
        if hasattr(tester, 'temp_dir') and tester.temp_dir.exists():
            shutil.rmtree(tester.temp_dir, ignore_errors=True)

if __name__ == "__main__":
    # 이벤트 루프 실행
    exit_code = asyncio.run(main())
    sys.exit(exit_code)